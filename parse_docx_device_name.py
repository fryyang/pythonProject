# coding=utf-8
from docx import Document
#读取文档
filename = "C:\\Users\\yanpe\\Desktop\\设备硬件信息统计表.docx";
doc = Document(filename) #filename为word文档

#获取文档中的表格
i = 0
for p in doc.paragraphs:
	style_name = p.style.name
	if style_name.startswith('Heading 4'):
		print(style_name, p.text, sep=':')
		i = i + 1
print(i)
#print(doc.paragraphs)
#读取第1个表格
#从4开始
#tb1=doc.tables[5]

#获取第一个表格的行
#tb1.rows  #获取表格的行数len(tb1.rows)
#print(len(tb1.rows))


