# coding=utf-8
from docx import Document
#读取文档
filename = "C:\\Users\\yanpe\\Desktop\\20211015.docx";
doc = Document(filename) #filename为word文档

#获取文档中的表格
doc.tables  #获取文档的表格个数 len(doc.tables)
print(len(doc.tables))
#读取第1个表格
#从4开始 61结束
j = 0
for i in range(4, len(doc.tables)-4):
	tb1=doc.tables[i]
	row_cells = tb1.rows[2].cells
	#print(len(row_cells))
	# 读取第一行所有单元格的内容
	for cell in row_cells:
		if cell.text.startswith('用途'):
			print(cell.text)
			j = j+1
print(j)
#获取第一个表格的行
#tb1.rows  #获取表格的行数len(tb1.rows)
#print(len(tb1.rows))
#读取表格的第一行的单元格
#用途
#对接AFC“超核”，承载监视中心对外连接AFC网络，大数据中心、TCC、如易行
#对接AFC“超核”，承载监视中心对外连接AFC网络，大数据中心、TCC、如易行
#对接AFC“超核”，承载监视中心对外连接AFC网络，大数据中心、TCC、如易行



