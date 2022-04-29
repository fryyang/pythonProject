# coding=utf-8
from docx import Document
import MySQLdb
import re

# 读取文档
filename = "C:\\Users\\yanpe\\Desktop\\20211015.docx"
doc = Document(filename)  # filename为word文档

# 获取文档中的表格
doc.tables  # 获取文档的表格个数 len(doc.tables)
print(len(doc.tables))


def time_format(str):
    if '/' in str:
        if str.count('/') == 1:
            str = str + '/01'
        return str
    if '年' in str:
        str = re.sub(r'年', "/", str)
    if '月' in str:
        str = re.sub(r'月', "/", str)
        if '日' in str:
            str = re.sub(r'日', "", str)
        else:
            str = str + '01'
    return str


# 读取第1个表格
# 从4开始 61结束

# 声明需要插入数据库的变量
# 所有字符先判断是否为 / 如果是则跳过
# device_base_info
# 设备用途 只取用途后下一行
device_use = None
# 设备名称
device_name = None
# 设备型号
device_type = None
# 硬件告警信息
hardware_alarm_info = None
# 隐患信息 在另一个表device_hidden_danger_info
device_base_info_id = None
danger_info = None
software_bug_info = None
# CPU型号
cpu_type = None
# 内存
memory = None
# RAID模式
raid_mode = None
# 网卡接口数量
network_card_interface = None
# 上连方式
network_connect_way = None
# 电源
power_source = None
# 设备规格
device_format = None
# 板卡类型
board_type = None
# 磁盘规格
disk_format = None
# 其他接口
other_interface = None
# 其他组件
other_component = None
# IOS
# 操作系统
operating_system = None
# 数据库
data_base = None
# 存储方式
store_way = None
# 启动方式
start_way = None
# 位置信息
# 设备位置
location_info = None
# 维护位置
maintain_location = None
# 上线时间
online_time = None
# 是否出保 需要判断之后转为数字 0/1
is_expired_protect = 1
# 备件情况 无备件/无备件 0/1 有备件再去读取备件详情
is_spare_parts = 0
# 备件详情
spare_parts_detail = None
# 管理方式
manage_type = None
# 管理IP
manage_ip = None
# 管理主机
manage_host = None
# 高可用性
high_availability = None
# 故障频次 只要数字
fault_frequency_year = 0
# 最后维护时间 有可能是/
last_maintain_data = None
#账号密码
user_password = None

host = 'localhost'
port = 3306
user = 'root'
passwd = 'root'
db = 'temp_parse_word'
conn = MySQLdb.connect(host, user, passwd, db)
cursor = conn.cursor()

splitStr = '/'
j = 0
mysqlId = 101
hiddenDangerInfoId = 1
for i in range(4, len(doc.tables) - 4):
    tb1 = doc.tables[i]
    rows = tb1.rows
    for row in rows:
        row_cells = row.cells
        for index in range(len(row_cells)):
            if row_cells[index].text.startswith('用途'):
                num = index + 1
                if row_cells[num].text is not None:
                    device_use = row_cells[num].text.replace(' ', '')
                    print('用途', device_use)
                break
            if row_cells[index].text.startswith('设备名称'):
                num = index + 1
                if row_cells[num].text is not None:
                    device_name = row_cells[num].text.replace('-', '')
                    print('设备名称', device_name)
                # 设备型号
                num = num + 2
                if row_cells[num].text is not None:
                    if splitStr != row_cells[num].text:
                        device_type = row_cells[num].text
                    else:
                        device_type = None
                    print('设备型号', device_type)
                break
            elif row_cells[index].text.startswith('硬件告警信息'):
                num = index + 1
                if row_cells[num].text is not None:
                    hardware_alarm_info = row_cells[num].text
                    print('硬件告警信息', hardware_alarm_info)
                # 隐患信息
                num = num + 1
                if row_cells[num].text is not None:
                    if row_cells[num].text.startswith('隐患信息'):
                        num = num + 1
                        danger_info = row_cells[num].text
                        print('隐患信息', danger_info)
                    elif row_cells[num].text.startswith('软件BUG信息'):
                        num = num + 1
                        software_bug_info = row_cells[num].text
                        print('软件BUG信息', software_bug_info)
                break
            elif row_cells[index].text.startswith('CPU型号'):
                num = index + 1
                if row_cells[num].text is not None:
                    cpu_type = row_cells[num].text.replace('-', '')
                    print('CPU型号', cpu_type)
                num = num + 2
                if row_cells[num].text is not None:
                    memory = row_cells[num].text.replace('-', '')
                    print('内存', memory)
                break
            elif row_cells[index].text.startswith('磁盘规格'):
                num = index + 1
                if row_cells[num].text is not None:
                    disk_format = row_cells[num].text.replace('-', '')
                    print('磁盘规格', disk_format)
                num = num + 2
                if row_cells[num].text is not None:
                    raid_mode = row_cells[num].text.replace('-', '')
                    print('RAID模式', raid_mode)
                break
            elif row_cells[index].text.startswith('网卡接口数量'):
                num = index + 1
                if row_cells[num].text is not None:
                    network_card_interface = row_cells[num].text.replace('-', '')
                    print('网卡接口数量', network_card_interface)
                num = num + 2
                if row_cells[num].text is not None:
                    network_connect_way = row_cells[num].text.replace('-', '')
                    print('上连方式', network_connect_way)
                break
            elif row_cells[index].text.startswith('电源'):
                num = index + 1
                if row_cells[num].text is not None:
                    power_source = row_cells[num].text.replace('-', '')
                    print('电源', power_source)
                # 设备规格
                num = num + 2
                if row_cells[num].text is not None:
                    device_format = row_cells[num].text.replace('-', '')
                    print('设备规格', device_format)
                break
            elif row_cells[index].text.startswith('板卡类型'):
                num = index + 1
                if row_cells[num].text is not None:
                    board_type = row_cells[num].text
                    print('板卡类型', board_type)
                # 磁盘规格
                num = num + 2
                if row_cells[num].text is not None:
                    disk_format = row_cells[num].text
                    print('磁盘规格', disk_format)
                break
            elif row_cells[index].text.startswith('其他接口'):
                num = index + 1
                if row_cells[num].text is not None:
                    other_interface = row_cells[num].text.replace('-', '')
                    print('其他接口', other_interface)
                num = num + 2
                if row_cells[num].text is not None:
                    if splitStr != row_cells[num].text:
                        other_component = row_cells[num].text.replace('-', '')
                    else:
                        other_component = None
                    print('其他组件', other_component)
                break
            elif row_cells[index].text.startswith('IOS'):
                num = index + 1
                if row_cells[num].text is not None:
                    operating_system = row_cells[num].text.replace('-', '')
                    print('操作系统', operating_system)
                break
            elif row_cells[index].text.startswith('操作系统'):
                num = index + 1
                if row_cells[num].text is not None:
                    operating_system = row_cells[num].text.replace('-', '')
                    print('操作系统', operating_system)
                num = num + 2
                if row_cells[num].text is not None:
                    if splitStr != row_cells[num].text:
                        data_base = row_cells[num].text
                    else:
                        data_base = None
                    print('数据库', data_base)
                break
            elif row_cells[index].text.startswith('存储方式'):
                num = index + 1
                if row_cells[num].text is not None:
                    if splitStr != row_cells[num].text:
                        store_way = row_cells[num].text.replace('-', '')
                    else:
                        store_way = None
                    print('存储方式', store_way)
                num = num + 2
                if row_cells[num].text is not None:
                    if splitStr != row_cells[num].text:
                        start_way = row_cells[num].text.replace('-', '')
                    else:
                        start_way = None
                    print('启动方式', start_way)
                break
            elif row_cells[index].text.startswith('位置信息') or row_cells[index].text.startswith('设备位置'):
                num = index + 1
                if row_cells[num].text is not None:
                    location_info = row_cells[num].text
                    print('位置信息', location_info)
                num = num + 2
                if row_cells[num].text is not None:
                    if splitStr != row_cells[num].text:
                        maintain_location = row_cells[num].text
                    else:
                        maintain_location = None
                    print('维护位置', maintain_location)
                break
            elif row_cells[index].text.startswith('上线时间'):
                num = index + 1
                if row_cells[num].text is not None and row_cells[num].text.find('-', 0, len(row_cells[num].text)) == -1:
                    online_time = time_format(row_cells[num].text)
                    print('上线时间', online_time)
                num = num + 2
                if row_cells[num].text is not None:
                    if splitStr != row_cells[num].text:
                        if '已' in row_cells[num].text:
                            is_expired_protect = 1
                        else:
                            is_expired_protect = 0
                    print('是否出保', is_expired_protect)
                break
            elif row_cells[index].text.startswith('备件情况'):
                num = index + 1
                if splitStr != row_cells[num].text:
                    if '有' in row_cells[num].text:
                        is_spare_parts = 1
                    else:
                        is_spare_parts = 0
                    print('是否备件', is_spare_parts)
                num = num + 2
                if row_cells[num].text is not None:
                    if splitStr != row_cells[num].text:
                        spare_parts_detail = row_cells[num].text.replace('-', '')
                    else:
                        spare_parts_detail = None
                    print('备件详情', spare_parts_detail)
                break
            elif row_cells[index].text.startswith('管理方式'):
                num = index + 1
                if row_cells[num].text is not None:
                    manage_type = row_cells[num].text.replace('-', '')
                    print('管理方式', manage_type)
                num = num + 2
                if row_cells[num].text is not None:
                    if splitStr != row_cells[num].text:
                        manage_ip = row_cells[num].text.replace('-', '')
                    else:
                        manage_ip = None
                    print('管理IP', manage_ip)
                break
            elif row_cells[index].text.startswith('管理主机'):
                num = index + 1
                if row_cells[num].text is not None:
                    manage_host = row_cells[num].text.replace('-', '')
                    print('管理主机', manage_host)
                num = num + 2
                if row_cells[num].text is not None:
                    if splitStr != row_cells[num].text:
                        high_availability = row_cells[num].text.replace('-', '')
                    else:
                        high_availability = None
                    print('高可用性', high_availability)
                break
            elif row_cells[index].text.startswith('故障频次'):
                num = index + 1
                if row_cells[num].text is not None and row_cells[num].text.find('-', 0, len(row_cells[num].text)) == -1:
                    fault_frequency_year = list(row_cells[num].text)[0]
                    print('故障频次', fault_frequency_year)
                # 最后维护时间
                num = num + 2
                if row_cells[num].text is not None:
                    if splitStr != row_cells[num].text and row_cells[num].text.find('-', 0,
                                                                                    len(row_cells[num].text)) == -1:
                        if row_cells[num].text.find(' ', 0, len(row_cells[num].text)) == -1:
                            last_maintain_data = time_format(row_cells[num].text)
                        else:
                            last_maintain_data = time_format(row_cells[num].text.split(' ')[0])
                    else:
                        last_maintain_data = None
                    print('最后维护时间', last_maintain_data)
                break
            elif row_cells[index].text.startswith('账号密码'):
                num = index + 1
                if row_cells[num].text is not None and row_cells[num].text.find('-', 0, len(row_cells[num].text)) == -1:
                    user_password = row_cells[num].text
                    print('账号密码', user_password)
    data = {
        'id': mysqlId,
        'device_use': device_use,
        'device_name': device_name,
        'device_type': device_type,
        'hardware_alarm_info': hardware_alarm_info,
        'software_bug_info': software_bug_info,
        'disk_format': disk_format,
        'raid_mode': raid_mode,
        'network_card_interface': network_card_interface,
        'network_connect_way': network_connect_way,
        'cpu_type': cpu_type,
        'memory': memory,
        'power_source': power_source,
        'device_format': device_format,
        'board_type': board_type,
        'other_interface': other_interface,
        'other_component': other_component,
        'operating_system': operating_system,
        'data_base': data_base,
        'store_way': store_way,
        'start_way': start_way,
        'location_info': location_info,
        'maintain_location': maintain_location,
        'online_time': online_time,
        'manage_type': manage_type,
        'manage_ip': manage_ip,
        'manage_host': manage_host,
        'is_expired_protect': is_expired_protect,
        'is_spare_parts': is_spare_parts,
        'spare_parts_detail': spare_parts_detail,
        'high_availability': high_availability,
        'fault_frequency_year': fault_frequency_year,
        'last_maintain_data': last_maintain_data,
        'user_password': user_password,
        'create_by': 'admin'
    }
    table = 'device_base_info'
    keys = ', '.join(data.keys())
    values = ', '.join(['%s'] * len(data))
    sql = 'INSERT INTO {table}({keys}) VALUES ({values})'.format(table=table, keys=keys, values=values)
    try:
        cursor.execute(sql, tuple(data.values()))
        print('Successful')
        conn.commit()
    except Exception as err:
        print('mysql error ', err)
        conn.rollback()

    data = {
        'id': hiddenDangerInfoId,
        'device_base_info_id': mysqlId,
        'danger_info': danger_info,
        'create_by': 'admin'
    }
    table = 'device_hidden_danger_info'
    keys = ', '.join(data.keys())
    values = ', '.join(['%s'] * len(data))
    sql = 'INSERT INTO {table}({keys}) VALUES ({values})'.format(table=table, keys=keys, values=values)
    try:
        cursor.execute(sql, tuple(data.values()))
        print('Successful')
        conn.commit()
    except Exception as err:
        print('mysql error ', err)
        conn.rollback()
    mysqlId = mysqlId + 1
    hiddenDangerInfoId = hiddenDangerInfoId + 1
cursor.close()
conn.close()
# 获取第一个表格的行
# tb1.rows  #获取表格的行数len(tb1.rows)
# print(len(tb1.rows))
# 读取表格的第一行的单元格
# 用途
# 对接AFC“超核”，承载监视中心对外连接AFC网络，大数据中心、TCC、如易行
# 对接AFC“超核”，承载监视中心对外连接AFC网络，大数据中心、TCC、如易行
# 对接AFC“超核”，承载监视中心对外连接AFC网络，大数据中心、TCC、如易行
